from docx import Document
import re

def analyze_docx(path):
    try:
        doc = Document(path)
        placeholders = set()
        pattern = re.compile(r'\{\{([^}]+)\}\}')
        
        for para in doc.paragraphs:
            matches = pattern.findall(para.text)
            placeholders.update(matches)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = pattern.findall(para.text)
                        placeholders.update(matches)
                        
        print(f"Found placeholders in {path}:")
        for p in sorted(placeholders):
            print(f"- {p}")
            
    except Exception as e:
        print(f"Error analyzing {path}: {e}")

if __name__ == "__main__":
    analyze_docx("TIMBRADO.docx")
