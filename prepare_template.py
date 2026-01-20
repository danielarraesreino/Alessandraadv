from docx import Document
import os

def prepare_template():
    base_path = "TIMBRADO.docx" # Original in root
    target_path = "src/core/templates/documents/TIMBRADO.docx"
    
    if not os.path.exists(base_path):
        print(f"Error: {base_path} not found.")
        return

    doc = Document(base_path)
    
    # Add standard legal header and placeholders
    doc.add_paragraph("")
    doc.add_paragraph("AO JUÍZO DA VARA CÍVEL DA COMARCA DE CAMPINAS/SP")
    doc.add_paragraph("")
    doc.add_paragraph("Processo nº: {{PROCESSO}}")
    doc.add_paragraph("")
    
    # Qualification
    p = doc.add_paragraph()
    p.add_run("{{CLIENTE}}").bold = True
    p.add_run(", inscrito no CPF/CNPJ sob nº {{CPF_CNPJ}}, vem, respeitosamente, à presença de Vossa Excelência, propor a presente")
    
    doc.add_paragraph("AÇÃO (NATUREZA A DEFINIR)", style='Heading 1')
    
    doc.add_paragraph("pelos fatos e fundamentos a seguir expostos.")
    doc.add_paragraph("")
    doc.add_paragraph("I - DOS FATOS")
    doc.add_paragraph("[Descrever os fatos aqui...]")
    doc.add_paragraph("")
    doc.add_paragraph("Campinas, {{DATA}}")
    doc.add_paragraph("")
    doc.add_paragraph("Alessandra M. Donadon")
    doc.add_paragraph("OAB/SP 123.456")

    doc.save(target_path)
    print(f"Template prepared at {target_path}")

if __name__ == "__main__":
    prepare_template()
