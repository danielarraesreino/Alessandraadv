import os
from docx import Document
from django.conf import settings
from django.utils import timezone
from apps.legal_cases.models import LegalCase

class DocumentAutomationService:
    """
    Service for generating legal documents (.docx) from templates.
    """
    
    TEMPLATES_DIR = os.path.join(settings.BASE_DIR, 'core', 'templates', 'documents')
    OUTPUT_DIR = os.path.join(settings.BASE_DIR, 'media', 'generated_docs')

    def __init__(self):
        if not os.path.exists(self.OUTPUT_DIR):
            os.makedirs(self.OUTPUT_DIR)

    def generate_base_document(self, legal_case: LegalCase, template_name: str = 'TIMBRADO.docx') -> str:
        """
        Generates a document by replacing placeholders in the official TIMBRADO.docx.
        
        Placeholders expected in .docx:
        {{CLIENTE}}, {{CPF_CNPJ}}, {{PROCESSO}}, {{DATA}}
        """
        template_path = os.path.join(self.TEMPLATES_DIR, template_name)
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template {template_name} not found in {self.TEMPLATES_DIR}")

        doc = Document(template_path)

        # Data map aligned with User Master Prompt
        data = {
            '{{CLIENTE}}': legal_case.client.full_name,
            '{{CPF_CNPJ}}': legal_case.client.cpf_cnpj,
            '{{PROCESSO}}': legal_case.process_number or "A definir",
            '{{DATA}}': timezone.now().strftime('%d/%m/%Y'),
            '{{AREA}}': legal_case.get_area_display()
        }

        # Replacement logic (Paragraphs)
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                    # Maintain formatting if possible
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, str(value))

        # Replacement logic (Tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in data.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, str(value))

        filename = f"peticao_{legal_case.id}_{timezone.now().strftime('%Y%m%d')}.docx"
        output_path = os.path.join(self.OUTPUT_DIR, filename)
        doc.save(output_path)
        
        return output_path
